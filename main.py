#!/usr/bin/env python3
"""
论文格式转换器 - 核心引擎
将 Markdown 论文稿按指定学术格式规范转换为排版规范的 DOCX 文档。

用法：
    python main.py --input paper.md --output paper.docx --format gbt7714
    python main.py --input paper.md --output paper.docx --format apa7 --custom-options custom.json
    python main.py --list-formats
"""

import argparse
import json
import os
import subprocess
import sys
import tempfile
from pathlib import Path

# 中文格式后处理（首行缩进、标题间距等）
try:
    from docx import Document
    from docx.shared import Pt, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

# ============================================================
# 论文格式规范 → md-to-docx 参数映射
# 数据来源：GB/T 7713.1-2006、GB/T 7714-2025、河北大学学位论文规范、
#           APA 7th、MLA 9th、Chicago Manual of Style、IMRaD
# ============================================================

FORMAT_PRESETS = {
    # ---- 中文学位论文 (GB/T 7713.1-2025 + GB/T 7714-2025 + 高校通用规范) ----
    # 格式依据：
    #   GB/T 7713.1-2025 §5.3.1: 每一章应另起页
    #   GB/T 7713.1-2025 §6.10: 天头≥25mm,订口≥25mm,地角≥20mm,切口≥20mm
    #   GB/T 7713.1-2025 §6.5: 章标题占2行,正文前空2字
    #   高校通用: 正文宋体小四(12pt)首行缩进2字符1.5倍行距两端对齐
    #   高校通用: 一级标题黑体三号(16pt)居中,二级黑体四号(14pt)左对齐,三级黑体小四(12pt)左对齐
    # ============================================================
    "gbt7714": {
        "display_name": "GB/T 7714-2025 学位论文格式（中文）",
        "description": "依据GB/T 7713.1-2025 + 高校通用排版规范：宋体小四正文、黑体标题、首行缩进2字符、1.5倍行距、每一章另起页",
        "options": {
            "documentType": "report",
            "style": {
                "fontFamily": "宋体",
                "paragraphSize": 24,           # 小四号 12pt (half-pts)
                "heading1Size": 32,             # 三号 16pt (half-pts)
                "heading2Size": 28,             # 四号 14pt (half-pts)
                "heading3Size": 24,             # 小四号 12pt (half-pts)
                "lineSpacing": 1.5,
                "headingSpacing": 12,
                "paragraphSpacing": 0,
                "paragraphAlignment": "JUSTIFIED",
                "headingAlignment": "LEFT",
            },
            "template": {
                "page": {
                    "margin": {
                        "top": 1417,            # 2.5cm → twips
                        "bottom": 1417,         # 2.5cm
                        "left": 1701,           # 3.0cm（装订侧）
                        "right": 1134,          # 2.0cm
                    }
                }
            },
        },
        "sections": [
            {"type": "cover", "title": "封面"},
            {"type": "abstract", "title": "摘要"},
            {"type": "abstract_en", "title": "Abstract"},
            {"type": "toc", "title": "目录"},
            {"type": "body", "title": "正文"},
            {"type": "references", "title": "参考文献"},
            {"type": "appendix", "title": "附录"},
            {"type": "acknowledgement", "title": "致谢"},
        ],
        "csl_style": "gb-t-7714-2025-numeric",
        "chinese_formatting": True,  # 启用中文格式后处理
    },

    # ---- APA 7th Edition ----
    "apa7": {
        "display_name": "APA 7th Edition",
        "description": "美国心理学会第7版格式，社会科学、教育学、心理学常用。标题页、作者-日期引用、参考文献悬挂缩进",
        "options": {
            "fontFamily": "Times New Roman",
            "paragraphSize": 12,
            "lineSpacing": 2.0,
            "pageMarginTop": 2.54,
            "pageMarginBottom": 2.54,
            "pageMarginLeft": 2.54,
            "pageMarginRight": 2.54,
        },
        "sections": [
            {"type": "title_page", "title": ""},
            {"type": "abstract", "title": "Abstract"},
            {"type": "body", "title": ""},
            {"type": "references", "title": "References"},
            {"type": "appendix", "title": "Appendix"},
        ],
    },

    # ---- MLA 9th Edition ----
    "mla9": {
        "display_name": "MLA 9th Edition",
        "description": "现代语言协会第9版格式，文学、语言、文化研究等人文学科常用。作者-页码引用、无独立标题页",
        "options": {
            "fontFamily": "Times New Roman",
            "paragraphSize": 12,
            "lineSpacing": 2.0,
            "pageMarginTop": 2.54,
            "pageMarginBottom": 2.54,
            "pageMarginLeft": 2.54,
            "pageMarginRight": 2.54,
        },
        "sections": [
            {"type": "body", "title": ""},
            {"type": "works_cited", "title": "Works Cited"},
        ],
    },

    # ---- Chicago Manual of Style (Notes-Bibliography) ----
    "chicago": {
        "display_name": "Chicago Manual of Style (Notes-Bibliography)",
        "description": "芝加哥格式注释-参考文献体系，历史学、艺术史、人类学常用。脚注/尾注 + 参考文献列表",
        "options": {
            "fontFamily": "Times New Roman",
            "paragraphSize": 12,
            "lineSpacing": 2.0,
            "pageMarginTop": 2.54,
            "pageMarginBottom": 2.54,
            "pageMarginLeft": 2.54,
            "pageMarginRight": 2.54,
        },
        "sections": [
            {"type": "title_page", "title": ""},
            {"type": "body", "title": ""},
            {"type": "bibliography", "title": "Bibliography"},
        ],
    },

    # ---- IMRaD 国际学术论文结构 ----
    "imrad": {
        "display_name": "IMRaD 国际学术论文结构",
        "description": "Introduction-Methods-Results-Discussion 标准科学论文结构。自然科学、医学、工程学投稿常用",
        "options": {
            "fontFamily": "Times New Roman",
            "paragraphSize": 11,
            "lineSpacing": 1.5,
            "pageMarginTop": 2.54,
            "pageMarginBottom": 2.54,
            "pageMarginLeft": 2.54,
            "pageMarginRight": 2.54,
        },
        "sections": [
            {"type": "title_page", "title": ""},
            {"type": "abstract", "title": "Abstract"},
            {"type": "introduction", "title": "Introduction"},
            {"type": "methods", "title": "Methods"},
            {"type": "results", "title": "Results"},
            {"type": "discussion", "title": "Discussion"},
            {"type": "conclusion", "title": "Conclusion"},
            {"type": "references", "title": "References"},
        ],
    },
}


def list_formats():
    """打印所有支持的格式规范"""
    print("支持的论文格式规范:\n")
    for key, preset in FORMAT_PRESETS.items():
        print(f"  {key:<15s}  {preset['display_name']}")
        print(f"  {'':15s}  {preset['description']}")
        print()


def build_options(format_name: str, custom_options: dict = None) -> dict:
    """根据格式名称构建 md-to-docx options JSON。

    从预设中加载排版参数，再与自定义 options 合并（自定义项覆写预设）。
    """
    if format_name not in FORMAT_PRESETS:
        available = ", ".join(FORMAT_PRESETS.keys())
        raise ValueError(f"未知格式 '{format_name}'。可用格式: {available}")

    preset = FORMAT_PRESETS[format_name]
    options = dict(preset["options"])

    if custom_options:
        options.update(custom_options)

    return options


def apply_chinese_formatting(docx_path: str, verbose: bool = False) -> None:
    """对 DOCX 执行中文论文格式标准后处理。

    依据 GB/T 7713.1-2025 + 高校通用规范 全面校正格式：
    - 正文：宋体小四(12pt)，1.5倍行距，首行缩进2字符，两端对齐
    - 一级标题(章)：黑体三号(16pt)，居中，段前24pt段后12pt，另起页
    - 二级标题(节)：黑体四号(14pt)，左对齐，段前12pt段后6pt
    - 三级标题(子节)：黑体小四(12pt)加粗，左对齐，段前6pt段后3pt
    - 英数：Times New Roman
    """
    if not HAS_PYTHON_DOCX:
        if verbose:
            print("[警告] python-docx 不可用，跳过中文格式后处理")
        return

    import re
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Emu

    doc = Document(docx_path)

    # ── 常量 ──
    # 字号（twips = 1/20pt）
    SIZE_H1 = Pt(16)   # 三号
    SIZE_H2 = Pt(14)   # 四号
    SIZE_H3 = Pt(12)   # 小四
    SIZE_BODY = Pt(12) # 小四

    # 段前段后（pt）
    SPACE_BEFORE_H1 = Pt(24)
    SPACE_AFTER_H1 = Pt(12)
    SPACE_BEFORE_H2 = Pt(12)
    SPACE_AFTER_H2 = Pt(6)
    SPACE_BEFORE_H3 = Pt(6)
    SPACE_AFTER_H3 = Pt(3)

    # 首行缩进2字符 = 2 * 12pt = 24pt = 480 twips
    INDENT_TWIPS = 480

    # ── 检测用正则 ──
    # H1 模式：第一章 一、二、1 2、第一节 等
    h1_pattern = re.compile(
        r'^('
        r'第[一二三四五六七八九十百千]+[章节篇条]|'  # 第一章
        r'[一二三四五六七八九十]+[、\s]|'            # 一、二、
        r'前言|引言|绪论|结语|结论|参考文献|附录|致谢'
        r')'
    )
    # H2 模式：1.1 2.1 等
    h2_pattern = re.compile(r'^\d+\.\d+\s')

    def set_run_font(run, ea_font: str, ascii_font: str = 'Times New Roman',
                     size: Pt = None, bold: bool = None):
        """设置 run 的中文/西文字体"""
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), ea_font)
        rFonts.set(qn('w:ascii'), ascii_font)
        rFonts.set(qn('w:hAnsi'), ascii_font)
        if size is not None:
            run.font.size = size
        if bold is not None:
            run.font.bold = bold

    def classify_paragraph(para, idx: int) -> str:
        """返回 heading1 / heading2 / heading3 / body"""
        text = para.text.strip()
        if not text:
            return 'body'

        # 测量最大字号和是否加粗
        max_size = 0
        any_bold = False
        for run in para.runs:
            if run.font.size and run.font.size > max_size:
                max_size = run.font.size
            if run.bold:
                any_bold = True

        # 文字前缀判断
        if h1_pattern.match(text):
            return 'heading1'
        if h2_pattern.match(text):
            return 'heading2'

        # 基于字号判断
        if max_size >= Emu(16 * 12700):  # ≥ 16pt
            return 'heading1'
        if max_size >= Emu(14 * 12700):  # ≥ 14pt
            return 'heading2'
        # 加粗 + 字号 ≥ 12pt 或 加粗 + 短文本（标题特征）
        if any_bold and max_size >= Emu(12 * 12700):
            return 'heading3'
        # 加粗短文本（无编号标题如"摘要""Abstract"）
        if any_bold and len(text) <= 30:
            return 'heading3'

        return 'body'

    total = len(doc.paragraphs)
    prev_was_h1 = False

    for idx, para in enumerate(doc.paragraphs):
        pf = para.paragraph_format
        text = para.text.strip()
        if not text:
            continue

        level = classify_paragraph(para, idx)

        if level == 'heading1':
            # 章标题：黑体三号(16pt)居中 段前24pt段后12pt 另起页
            for run in para.runs:
                set_run_font(run, '黑体', 'Times New Roman', SIZE_H1, bold=True)
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = None
            pf.space_before = SPACE_BEFORE_H1
            pf.space_after = SPACE_AFTER_H1
            # 另起页（第一段不加分页符）
            if idx > 0:
                pf.page_break_before = True
            prev_was_h1 = True

        elif level == 'heading2':
            # 节标题：黑体四号(14pt)左对齐 段前12pt段后6pt
            for run in para.runs:
                set_run_font(run, '黑体', 'Times New Roman', SIZE_H2, bold=True)
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = None
            pf.space_before = SPACE_BEFORE_H2
            pf.space_after = SPACE_AFTER_H2
            prev_was_h1 = False

        elif level == 'heading3':
            # 子节：黑体小四(12pt)加粗左对齐 段前6pt段后3pt
            for run in para.runs:
                set_run_font(run, '黑体', 'Times New Roman', SIZE_H3, bold=True)
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = None
            pf.space_before = SPACE_BEFORE_H3
            pf.space_after = SPACE_AFTER_H3
            prev_was_h1 = False

        else:
            # 正文：宋体小四(12pt) 首行缩进2字符 两端对齐 1.5倍行距
            for run in para.runs:
                set_run_font(run, '宋体', 'Times New Roman', SIZE_BODY)
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.first_line_indent = Twips(INDENT_TWIPS)
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            prev_was_h1 = False

    doc.save(docx_path)
    if verbose:
        h1_count = sum(1 for p in doc.paragraphs if classify_paragraph(p, 0) == 'heading1')
        h2_count = sum(1 for p in doc.paragraphs if classify_paragraph(p, 0) == 'heading2')
        h3_count = sum(1 for p in doc.paragraphs if classify_paragraph(p, 0) == 'heading3')
        print(f"[后处理] 中文论文格式已应用："
              f"H1×{h1_count} H2×{h2_count} H3×{h3_count} 正文×{total - h1_count - h2_count - h3_count}")
        print(f"         正文: 宋体小四 1.5倍行距 首行缩进2字符 两端对齐")
        print(f"         一级标题: 黑体三号 居中 段前24pt段后12pt 另起页")
        print(f"         二级标题: 黑体四号 左对齐 段前12pt段后6pt")
        print(f"         三级标题: 黑体小四加粗 左对齐 段前6pt段后3pt")


def run_md_to_docx(input_path: str, output_path: str, options: dict, verbose: bool = False) -> str:
    """调用 npx @mohtasham/md-to-docx 执行转换。

    Args:
        input_path: 输入 Markdown 文件路径
        output_path: 输出 DOCX 文件路径
        options: 排版参数 dict
        verbose: 是否输出详细日志

    Returns:
        生成的 DOCX 文件路径

    Raises:
        RuntimeError: 转换失败
        subprocess.TimeoutExpired: 超时
    """
    # 将 options 写入临时 JSON 文件
    tmp_dir = tempfile.mkdtemp()
    options_path = os.path.join(tmp_dir, "options.json")

    with open(options_path, "w", encoding="utf-8") as f:
        json.dump(options, f, ensure_ascii=False, indent=2)

    if verbose:
        print(f"[配置] 排版参数:\n{json.dumps(options, ensure_ascii=False, indent=2)}")

    cmd = [
        "npx", "@mohtasham/md-to-docx",
        input_path,
        output_path,
        "--options", options_path,
    ]

    if verbose:
        print(f"[执行] {' '.join(cmd)}")

    result = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        timeout=120,
    )

    if result.returncode != 0:
        error_msg = (result.stderr or result.stdout or "未知错误").strip()
        # npx 安装日志不算错误
        if " installed" in error_msg and result.returncode == 0:
            pass
        else:
            raise RuntimeError(
                f"md-to-docx 转换失败 (code={result.returncode}): {error_msg}"
            )

    if verbose and result.stdout:
        print(f"[输出] {result.stdout.strip()}")

    # 验证输出文件
    output_file = Path(output_path)
    if not output_file.exists():
        raise RuntimeError(f"转换过程完成但输出文件未找到: {output_path}")

    return str(output_file.resolve())


def main():
    parser = argparse.ArgumentParser(
        description="论文格式转换器 — Markdown → 规范 DOCX",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例:\n"
            "  python main.py -i paper.md -o paper.docx -f gbt7714\n"
            "  python main.py -i paper.md -o paper.docx -f apa7 -c custom.json\n"
            "  python main.py --list-formats\n"
        ),
    )
    parser.add_argument("--input", "-i", required=True, help="输入 Markdown 文件路径")
    parser.add_argument("--output", "-o", required=True, help="输出 DOCX 文件路径")
    parser.add_argument(
        "--format", "-f",
        default="gbt7714",
        help=f"格式规范（默认: gbt7714）。可用: {', '.join(FORMAT_PRESETS.keys())}",
    )
    parser.add_argument(
        "--custom-options", "-c",
        help="自定义 JSON options 文件路径，可覆写预设中的任何参数",
    )
    parser.add_argument("--verbose", "-v", action="store_true", help="详细输出日志")
    parser.add_argument("--list-formats", action="store_true", help="列出所有支持的格式规范")

    args = parser.parse_args()

    # 列出格式
    if args.list_formats:
        list_formats()
        return

    # 验证输入文件
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"错误: 输入文件不存在: {args.input}", file=sys.stderr)
        sys.exit(1)

    # 加载自定义 options
    custom_options = None
    if args.custom_options:
        custom_path = Path(args.custom_options)
        if not custom_path.exists():
            print(f"错误: 自定义 options 文件不存在: {args.custom_options}", file=sys.stderr)
            sys.exit(1)
        try:
            with open(custom_path, "r", encoding="utf-8") as f:
                custom_options = json.load(f)
        except json.JSONDecodeError as e:
            print(f"错误: 自定义 options JSON 解析失败: {e}", file=sys.stderr)
            sys.exit(1)

    try:
        options = build_options(args.format, custom_options)
        result_path = run_md_to_docx(
            str(input_path.resolve()),
            str(Path(args.output).resolve()),
            options,
            args.verbose,
        )

        preset = FORMAT_PRESETS.get(args.format, {})
        format_name = preset.get("display_name", args.format)

        # 中文格式后处理（首行缩进等）
        if preset.get("chinese_formatting"):
            try:
                apply_chinese_formatting(result_path, args.verbose)
                print(f"   中文格式: 首行缩进2字符、两端对齐 ✅")
            except Exception as e:
                print(f"   ⚠️ 中文格式后处理异常: {e}")

        print(f"✅ 转换完成")
        print(f"   输出: {result_path}")
        print(f"   格式: {format_name}")

        # 摘要关键参数
        style = options.get("style", {})
        key_params = {k: style[k] for k in
                      ["fontFamily", "paragraphSize", "lineSpacing"]
                      if k in style}
        if key_params:
            print(f"   参数: {json.dumps(key_params, ensure_ascii=False)}")

    except ValueError as e:
        print(f"❌ 参数错误: {e}", file=sys.stderr)
        sys.exit(1)
    except subprocess.TimeoutExpired:
        print("❌ 转换超时（120秒），请检查网络或输入文件大小", file=sys.stderr)
        sys.exit(1)
    except RuntimeError as e:
        print(f"❌ 转换失败: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
